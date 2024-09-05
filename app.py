import traceback
import torch
import flask
from flask import Flask,render_template,request,redirect,url_for,session,flash
from werkzeug.utils import secure_filename
import os
import json
from urllib.request import urlopen
import bs4 as bs
from flask import request
from docx import Document
import PyPDF2
import requests
from rouge import Rouge
from transformers import BartTokenizer, BartForConditionalGeneration, BartConfig
BART_PATH = 'bart_finetuned'

app=Flask(__name__,template_folder='templates',static_folder="static")

model = BartForConditionalGeneration.from_pretrained(BART_PATH, output_past=True)
tokenizer = BartTokenizer.from_pretrained(BART_PATH, output_past=True)
#device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')

# Define the directory where uploaded files will be stored
def generate_summary(text):
    input_ids = tokenizer.encode(text, return_tensors='pt', max_length=2500, truncation=True)
    summary_ids = model.generate(input_ids, num_beams=4, min_length=50, max_length=2500, early_stopping=True)
    summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
    return summary 
def summarize_with_bart(text):
    # Split the input text into smaller chunks
    chunks = [text[i:i + 1000] for i in range(0, len(text), 1000)]

    # Summarize each chunk
    summaries = []
    for chunk in chunks:
        input_ids = tokenizer.encode(chunk, return_tensors='pt', max_length=2500, truncation=True)
        summary_ids = model.generate(input_ids, num_beams=4, min_length=50, max_length=2500, early_stopping=True)
        summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
        summaries.append(summary)

    # Combine individual summaries
    overall_summary = ' '.join(summaries)
    return overall_summary
@app.route('/summarize_file', methods=['POST'])
def summarize_file():
    UPLOAD_FOLDER = 'uploads'  # Define the directory where uploaded files will be stored
    app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER  # Configure Flask app with the upload folder

    if not os.path.exists(UPLOAD_FOLDER):  # Create the directory if it doesn't exist
        os.makedirs(UPLOAD_FOLDER)
    uploaded_file = request.files['file']
    if uploaded_file.filename != '':
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(uploaded_file.filename))
        uploaded_file.save(file_path)
        try:
            if uploaded_file.filename.endswith('.pdf'):
                text = extract_text_from_pdf(uploaded_file)
                summary = generate_summary(text)
                return render_template('uploads.html', summary=summary)
            else:
                if uploaded_file.filename.endswith('.txt'):
                    text = extract_text_from_txt(file_path)
                elif uploaded_file.filename.endswith('.docx'):
                    # Save the uploaded file to a temporary location
                    # Pass the file path to the extract_text_from_docx function
                    text = extract_text_from_docx(file_path)
                    # Delete the temporary file
                    os.remove(file_path)
                else:
                    return "Unsupported file format"
                summary = summarize_with_bart(text)
                return render_template('uploads.html', summary=summary)
            
        except Exception as e:
            return f"Error processing the file: {e}"
        except Exception as e:
            traceback.print_exc()  # Print the full traceback
        return f"Error processing the file: {e}"

    else:
                return "No file provided"

def extract_text_from_pdf(file):
    text = ""
    pdf_reader = PyPDF2.PdfReader(file)
    num_pages = len(pdf_reader.pages)
    for page_number in range(num_pages):
        page = pdf_reader.pages(page_number)
        text += page.extractText()
    return text
def extract_text_from_docx(file):
    try:
        if os.path.exists(file):
            doc = Document(file)
            text = ""
            if doc.paragraphs:
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"
                return text
            else:
                return "No text found in the document."
        else:
            return f"File not found at '{file}'"
    except Exception as e:
        return f"Error extracting text from DOCX: {e}"
def extract_text_from_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as txt_file:
            text = txt_file.read()
        return text
    except Exception as e:
        return f"Error extracting text from TXT: {e}"
#def calculate_rouge(hypotheses, references):
    #rouge = Rouge()
    #scores = rouge.get_scores(hypotheses, references)
    #return scores
@app.route('/')
def  home():
    return render_template('index.html')
@app.route('/summarize')
def summarize():
   return render_template('summarize.html')
@app.route('/summarize', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        input_text = request.form['input_text']
        summary = generate_summary(input_text)
        
        return render_template('summarize.html', input_text=input_text, output_text=summary)
    return render_template('index.html', input_text='', output_text='')

@app.route('/summarize/uploads')
def uploads():
   return render_template('uploads.html')
@app.route('/summarize/web')
def web():
   return render_template('url.html')
@app.route('/contact')
def contact():
   return render_template('contact.html')
@app.route('/about')
def about():
   return  render_template('about.html')
def get_only_text(url):
   try:
        response = requests.get(url)
        if response.status_code == 200:
            content_type = response.headers.get('content-type')
            if 'application/pdf' in content_type:
                return extract_text_from_pdf(response.content)
            elif 'text/html' in content_type:
                soup = bs.BeautifulSoup(response.text, 'html.parser')
                title = soup.title.text
                text = ''.join(map(lambda p: p.text, soup.find_all('p')))
                return title, text
            else:
                return "Error", "Unsupported content type: " + content_type
        else:
            return "Error", "Unable to fetch content from URL"
   except Exception as e:
        return "Error", str(e)
        return "Error", str(e)
@app.route('/get_text', methods=['POST'])
def get_text():
    if request.method == 'POST':
        url = request.form['url']
        title, text = get_only_text(url)
        summary = summarize_with_bart(text)
        return render_template('result.html', title=title, text=text, summary=summary)

if __name__=="__main__":
   
   app.run(debug=True)
